"""Convert RD_IMPLEMENTATION_PLAN.md → RD_Implementation_Plan.docx.

Focused markdown subset supported (matches what the source doc uses):
- ATX headings (# … ######) → Word Heading 1-6
- Paragraphs separated by blank lines
- Pipe tables with a header separator row → Word tables
- Unordered lists (-, *) and ordered lists (1.) → bulleted / numbered paragraphs
- Fenced code blocks (```) → monospaced shaded paragraphs
- Inline: **bold**, `code`, [text](url)
- Horizontal rule (---) → page break
- > blockquote (rare; rendered indented)

Run: python3 build_rd_plan_docx.py
"""
from __future__ import annotations
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Cm

SRC = Path(__file__).parent / 'RD_IMPLEMENTATION_PLAN.md'
DST = Path(__file__).parent / 'RD_Implementation_Plan.docx'

BRAND_CYAN = RGBColor(0x2E, 0x8B, 0xAD)
SLATE_DEEP = RGBColor(0x3D, 0x46, 0x4D)
SLATE = RGBColor(0x5A, 0x65, 0x6F)
GRAY_BG = 'F4F6F8'

# ─── Helpers ──────────────────────────────────────────────────────────────────

INLINE_BOLD = re.compile(r'\*\*(.+?)\*\*')
INLINE_CODE = re.compile(r'`([^`]+)`')
INLINE_LINK = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')

def _shade_cell(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)


def _add_inline(paragraph, text, *, base_font=None, base_size=None, color=None):
    """Render a markdown line with **bold**, `code`, [link](url) into runs."""
    # Tokenize order: links first, then bold, then code. Use a single pass with
    # alternating split.
    pos = 0
    pattern = re.compile(
        r'(\*\*.+?\*\*)|(`[^`]+`)|(\[[^\]]+\]\([^)]+\))'
    )
    for m in pattern.finditer(text):
        if m.start() > pos:
            _emit_run(paragraph, text[pos:m.start()],
                      base_font=base_font, base_size=base_size, color=color)
        tok = m.group(0)
        if tok.startswith('**'):
            _emit_run(paragraph, tok[2:-2], bold=True,
                      base_font=base_font, base_size=base_size, color=color)
        elif tok.startswith('`'):
            _emit_run(paragraph, tok[1:-1], code=True,
                      base_size=base_size, color=color)
        else:  # link
            link_m = INLINE_LINK.match(tok)
            label, url = link_m.group(1), link_m.group(2)
            _add_hyperlink(paragraph, label, url, base_size=base_size)
        pos = m.end()
    if pos < len(text):
        _emit_run(paragraph, text[pos:],
                  base_font=base_font, base_size=base_size, color=color)


def _emit_run(paragraph, txt, *, bold=False, italic=False, code=False,
              base_font=None, base_size=None, color=None):
    if not txt:
        return
    run = paragraph.add_run(txt)
    if bold:   run.bold = True
    if italic: run.italic = True
    if code:
        run.font.name = 'Menlo'
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
            rFonts.set(qn(attr), 'Menlo')
        rPr.append(rFonts)
        # Light shading on inline code
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), GRAY_BG)
        rPr.append(shd)
    elif base_font:
        run.font.name = base_font
    if base_size: run.font.size = base_size
    if color:     run.font.color.rgb = color


def _add_hyperlink(paragraph, text, url, *, base_size=None):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color'); color.set(qn('w:val'), '2E8BAD')
    underline = OxmlElement('w:u'); underline.set(qn('w:val'), 'single')
    rPr.append(color); rPr.append(underline)
    if base_size:
        sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(base_size.pt * 2)))
        rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text; t.set(qn('xml:space'), 'preserve')
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# ─── Document setup ───────────────────────────────────────────────────────────

def _configure_styles(doc):
    """Brand the headings + body styles."""
    styles = doc.styles
    body = styles['Normal']
    body.font.name = 'Calibri'
    body.font.size = Pt(11)
    body.paragraph_format.space_after = Pt(6)
    body.paragraph_format.line_spacing = 1.25

    for level, size in [(1, 22), (2, 16), (3, 13), (4, 12)]:
        s = styles[f'Heading {level}']
        s.font.name = 'Calibri'
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = SLATE_DEEP if level > 1 else BRAND_CYAN
        s.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.keep_with_next = True


# ─── Block-level conversion ───────────────────────────────────────────────────

def _flush_paragraph(doc, lines):
    if not lines:
        return
    p = doc.add_paragraph()
    text = ' '.join(line.strip() for line in lines)
    _add_inline(p, text)


def _flush_list(doc, items, ordered):
    style = 'List Number' if ordered else 'List Bullet'
    for txt in items:
        p = doc.add_paragraph(style=style)
        _add_inline(p, txt)


def _flush_code(doc, code_lines, lang):
    label = f'[{lang}]' if lang else '[code]'
    note = doc.add_paragraph()
    note_run = note.add_run(f'{label} block (rendered as plain text — see source MD for full visual)')
    note_run.italic = True
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = SLATE
    for line in code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Menlo'
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
            rFonts.set(qn(attr), 'Menlo')
        rPr.append(rFonts)
        run.font.size = Pt(9)
        run.font.color.rgb = SLATE_DEEP
    doc.add_paragraph()  # trailing space


def _flush_table(doc, rows):
    """rows: list of list-of-strings; first row is header."""
    if not rows: return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            txt = row[c_idx] if c_idx < len(row) else ''
            cell.text = ''
            p = cell.paragraphs[0]
            _add_inline(p, txt, base_size=Pt(10))
            if r_idx == 0:
                _shade_cell(cell, '2E8BAD')
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    doc.add_paragraph()


# ─── Main parser ──────────────────────────────────────────────────────────────

H_RE = re.compile(r'^(#{1,6})\s+(.*)$')
TABLE_SEP_RE = re.compile(r'^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$')
LIST_RE = re.compile(r'^(\s*)([-*])\s+(.*)$')
OLIST_RE = re.compile(r'^(\s*)\d+\.\s+(.*)$')
HR_RE = re.compile(r'^\s*---+\s*$')
CODE_FENCE_RE = re.compile(r'^```(\w*)\s*$')


def convert(md_text: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    _configure_styles(doc)

    lines = md_text.split('\n')
    i = 0
    paragraph_buf: list[str] = []
    list_buf: list[str] = []
    list_ordered = False
    table_buf: list[list[str]] = []
    code_buf: list[str] = []
    in_code = False
    code_lang = ''

    def flush_text():
        nonlocal paragraph_buf, list_buf
        if paragraph_buf:
            _flush_paragraph(doc, paragraph_buf)
            paragraph_buf = []
        if list_buf:
            _flush_list(doc, list_buf, list_ordered)
            list_buf = []

    while i < len(lines):
        line = lines[i]
        # Code fence toggle
        m_code = CODE_FENCE_RE.match(line)
        if m_code:
            if in_code:
                _flush_code(doc, code_buf, code_lang)
                code_buf = []; in_code = False; code_lang = ''
            else:
                flush_text()
                in_code = True; code_lang = m_code.group(1)
            i += 1; continue
        if in_code:
            code_buf.append(line); i += 1; continue

        # Headings
        h = H_RE.match(line)
        if h:
            flush_text()
            level = min(len(h.group(1)), 4)
            text = h.group(2).strip()
            p = doc.add_paragraph(style=f'Heading {level}')
            _add_inline(p, text)
            i += 1; continue

        # Horizontal rule → page break
        if HR_RE.match(line):
            flush_text()
            p = doc.add_paragraph()
            p.add_run().add_break(WD_BREAK.PAGE)
            i += 1; continue

        # Tables: detect a header row immediately followed by a separator row
        if line.lstrip().startswith('|') and i + 1 < len(lines) and TABLE_SEP_RE.match(lines[i+1]):
            flush_text()
            # Collect header
            def split_row(s):
                s = s.strip().strip('|')
                # split on pipes that aren't escaped (no escaping in our doc)
                return [c.strip() for c in s.split('|')]
            header = split_row(line)
            i += 2  # skip header + separator
            rows = [header]
            while i < len(lines) and lines[i].lstrip().startswith('|'):
                rows.append(split_row(lines[i]))
                i += 1
            _flush_table(doc, rows)
            continue

        # Lists
        m_ul = LIST_RE.match(line)
        m_ol = OLIST_RE.match(line)
        if m_ul or m_ol:
            if paragraph_buf: _flush_paragraph(doc, paragraph_buf); paragraph_buf = []
            if not list_buf:
                list_ordered = bool(m_ol)
            if m_ul:
                list_buf.append(m_ul.group(3))
            else:
                list_buf.append(m_ol.group(2))
            i += 1; continue

        # Blockquote
        if line.startswith('> '):
            flush_text()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run_pre = p.add_run('▸ ')
            run_pre.font.color.rgb = BRAND_CYAN
            _add_inline(p, line[2:].strip())
            i += 1; continue

        # Blank line ends paragraph / list
        if not line.strip():
            flush_text()
            i += 1; continue

        paragraph_buf.append(line)
        i += 1

    # End of document
    if in_code:
        _flush_code(doc, code_buf, code_lang)
    flush_text()
    return doc


def main():
    md = SRC.read_text(encoding='utf-8')
    doc = convert(md)
    doc.save(DST)
    print(f'Wrote {DST} ({DST.stat().st_size:,} bytes)')


if __name__ == '__main__':
    main()
